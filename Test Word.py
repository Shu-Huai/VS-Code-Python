from docx import Document

docxFile = Document()
docxFile.add_heading("这是一个一级标题", level=1)
docxFile.add_paragraph("这是一个副级标题", "Title")
A = docxFile.add_paragraph("My name is aaa")
A.add_run("我学习的很快乐，啊哈哈哈哈哈，非常好 Good!!!")
docxFile.add_heading("这是一个二级标题", level=2)
A = docxFile.add_paragraph("这个是二级标题的内容呀")
B = A.add_run("二级标题里面的正文 继续添加！！！！！！！")
B.font.bold = True
B.font.size = (20)
docxFile.add_heading("我爱学习Python以下就是python的logo呀", level=3)
docxFile.add_table(rows=5, cols=5)
docxFile.save("Test Word.docx")