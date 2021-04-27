import docx
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def GetMaxPage():
    url = []
    url.append("https://store.steampowered.com/search/?specials=1&page=1")
    soup = GetSoup(GetContent(url))
    node = soup[0].find_all("div", class_="search_pagination_right")
    return int(node[0].contents[5].contents[0])


def GetUrls(pages):
    urls = []
    for i in range(pages):
        urlExample = "https://store.steampowered.com/search/?specials=1&page={}".format(i + 1)
        urls.append(urlExample)
    return urls


def GetContent(urls):
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.116 Safari/537.36'}
    responseList = []
    contentList = []
    for i in range(len(urls)):
        try:
            responseList.append(requests.get(urls[i], headers=headers))
        except ConnectionError:
            break
        contentList.append(responseList[i].text)
    return contentList


def GetSoup(contentList):
    soup = []
    for i in range(len(contentList)):
        soup.append(BeautifulSoup(contentList[i], "html.parser"))
    return soup


def GetGameName(contentList):
    names = []
    soup = GetSoup(contentList)
    for i in range(len(contentList)):
        names.extend(soup[i].find_all("span", class_="title"))
    for i in range(len(names)):
        names[i] = names[i].string
    return names


def GetGameUrl(contentList):
    urls = []
    soup = GetSoup(contentList)
    urlPrefix = "https://store.steampowered.com/"
    for i in range(len(contentList)):
        for node in soup[i].find_all("a"):
            temp = node.get("href")
            if (urlPrefix + "app/" in temp or urlPrefix + "bundle/" in temp or urlPrefix + "sub/" in temp) and "view" not in temp:
                urls.append(node.get("href"))
    return urls


def GetPrice(contentList):
    previousPrices = []
    nowPrices = []
    discounts = []
    soup = GetSoup(contentList)
    for i in range(len(contentList)):
        count = 0
        unpurchaseableIndex = []
        for node in soup[i].find_all("div", class_="col search_discount responsive_secondrow"):
            discount = node.text.strip("\n")
            if discount == "":
                discounts.append("0")
                unpurchaseableIndex.append(count)
            else:
                discounts.append(discount)
            count += 1
        for node in soup[i].find_all("div", class_="col search_price discounted responsive_secondrow"):
            previousPrices.append(node.contents[1].contents[0].contents[0])
            nowPrices.append(node.contents[3].strip())
        for j in unpurchaseableIndex:
            previousPrices.insert(i * 25 + j, "Unpurchasable")
            nowPrices.insert(i * 25 + j, "Unpurchasable")
    return previousPrices, nowPrices, discounts


def Merge(names, urls, previousPrices, nowPrices, discounts):
    games = []
    for i in range(len(names)):
        games.append(dict(gameName=names[i], gameUrl=urls[i], previousPrice=previousPrices[i], nowPrice=nowPrices[i], discount=discounts[i]))
    return games


def Sort(games):
    for i in range(len(games)):
        games[i]["discount"] = int(games[i]["discount"].strip("%"))
    games = sorted(games, key=lambda x: (x["discount"], x['gameName']))
    for i in range(len(games)):
        games[i]["discount"] = str(games[i]["discount"]) + "%"
    return games


def SaveToDocx(games):
    docxFile = Document()
    docxFile.styles["Normal"].font.name = "Times New Roman"
    docxFile.styles["Normal"].font.size = Pt(12)
    docxFile.add_paragraph("Here is the Steam discount information for this week.")
    for i in range(len(games)):
        paragraph = docxFile.add_paragraph()
        paragraph.add_run("Game: %s.\nLink: " % games[i]["gameName"])
        part = paragraph.part
        ralationId = part.relate_to(games[i]["gameUrl"], docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperLink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperLink.set(
            docx.oxml.shared.qn("r:id"),
            ralationId,
        )
        r = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        r.append(rPr)
        r.text = games[i]["gameUrl"]
        hyperLink.append(r)
        run = paragraph.add_run()
        run._r.append(hyperLink)
        run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        run.font.underline = True
        paragraph.add_run("\nDiscount: %s, Price: %s, Previous Price: %s.\n" % (games[i]["discount"], games[i]["nowPrice"], games[i]["previousPrice"]))
        docxFile.add_picture(r"Steam Discount Information Getter\Game Cover\Darkest Dungeon®.png")
    docxFile.save(r"Steam Discount Information Getter\Steam Discount Information.docx")


pages = input("Please input the pages you want, min is 1, max is %d, default is 5: " % GetMaxPage())
try:
    pages = int(pages)
except ValueError:
    pages = 5
urls = GetUrls(pages)
contentList = GetContent(urls)
gameNames = GetGameName(contentList)
gameUrls = GetGameUrl(contentList)
previousPrices, nowPrices, discounts = GetPrice(contentList)
games = Merge(gameNames, gameUrls, previousPrices, nowPrices, discounts)
games = Sort(games)
for i in range(len(games)):
    print("Game: %s.\nLink: %s." % (games[i]["gameName"], games[i]["gameUrl"]))
    print("Discount: %s, Price: %s, Previous Price: %s.\n" % (games[i]["discount"], games[i]["nowPrice"], games[i]["previousPrice"]))
filename = "Steam Discount Getter.md"
SaveToDocx(games)